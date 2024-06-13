from docx import Document

# Create a new Document
doc = Document()

# Add title
doc.add_heading('Course Design Manual', 0)

# Add Course Design Information
doc.add_heading('Course Design Information', level=1)
doc.add_paragraph("Course Design Name: Course Design for Principles of Concrete Structure Design")
doc.add_paragraph("Course Design Topic: Design of reinforced concrete cantilever beam")
doc.add_paragraph("College Name: [Your College Name]")
doc.add_paragraph("Professional Class: [Your Professional Class]")
doc.add_paragraph("Student ID: [Your Student ID]")
doc.add_paragraph("Student Name: [Your Name]")
doc.add_paragraph("Student Achievement: [Your Achievement]")
doc.add_paragraph("Mentor: [Your Mentor's Name]")
doc.add_paragraph("Course Design Time: 2023.12.04 to 2023.12.18")

# Add Design Basis
doc.add_heading('Design Basis', level=1)
doc.add_paragraph("Design Codes and References:")
doc.add_paragraph("1. Ministry of Housing and Urban-Rural Construction of the People's Republic of China. Code for Design of Concrete Structures GB50010—2010. Beijing: China Architecture & Building Press 2010.")
doc.add_paragraph("2. Southeast University, Tianjin University & Tongii University. The design theory for concrete structures. The fifth edition. Beijing: China Architecture & Building Press 2012.")
doc.add_paragraph("3. Jun Zhao, Xinling Wang, Liusheng Chu, Hui Qian, Le Li. Reinforced concrete fundamentals. Beijing: China Architecture & Building Press 2015.")
doc.add_paragraph("4. Xianglin Gu, Xianyu Jin, Yong Zhou. Basic principles of concrete structures. Shanghai: Tongji University Press 2015.")
doc.add_paragraph("5. Xingwen Liang, Qingxuan Shi. The design theory for concrete structures. The third edition. Beijing: China Architecture & Building Press 2016.")

# Add Design Steps
doc.add_heading('Design Steps', level=1)

# 2.1 Collecting Data and Finding the b x h
doc.add_heading('2.1 Collecting Data and Finding the b x h', level=2)
doc.add_paragraph(r"C30, f_ck = 20.1, f_yk = 500, E = 30000 N/mm^2")
doc.add_paragraph(r"l_0 = 8.0 m, d = 0.5 m, b = 0.3 m")
doc.add_paragraph("Determination of Section Dimensions:")
doc.add_paragraph(r"h = l_0 / 16 = 8 / 16 = 0.5 m")
doc.add_paragraph(r"b = h / 1.5 = 0.5 / 1.5 = 0.33 m")
doc.add_paragraph("Calculation of Steel Reinforcement:")
doc.add_paragraph(r"A_s = M / (f_yk \cdot z)")
doc.add_paragraph(r"= 1.25 \cdot 300 / (500 \cdot 0.85 \cdot 0.5)")
doc.add_paragraph(r"= 0.014 m^2")

# 2.2 Calculating the Sectional Area of the Longitudinal Reinforcement
doc.add_heading('2.2 Calculating the Sectional Area of the Longitudinal Reinforcement', level=2)
doc.add_paragraph(r"V_u = V_cs")
doc.add_paragraph(r"V_cs = \alpha f_{ck} b_{w} d + f_{yk} \frac{A_{s}}{s} h_{0}")
doc.add_paragraph(r"= 0.7 \cdot 420 \cdot 833.018 \cdot 793.3 + 300 \cdot \frac{200 \cdot 200}{200}")
doc.add_paragraph(r"= 71825772.30")
doc.add_paragraph(r"V_s0 = 0.8 f_{yk} A_{s} \sin \alpha")
doc.add_paragraph(r"= 0.8 \cdot 360 \cdot 968.2 \cdot \sin 45")
doc.add_paragraph(r"= 164410.8")
doc.add_paragraph(r"V_u = \alpha f_{ck} b_{w} d + 0.8 f_{yk} A_{s} \sin \alpha")
doc.add_paragraph(r"= 71825772.30 + 164410.8")
doc.add_paragraph(r"= 71990183.1")

# 3.1 Calculating the Stirrups
doc.add_heading('3.1 Calculating the Stirrups', level=2)
doc.add_paragraph("Calculation of Moment Capacity:")
doc.add_paragraph(r"M_u = R_B \cdot 6 - q_1 \cdot (6 / 2)^2")
doc.add_paragraph(r"= 530.00 \cdot 6 - 60 \cdot 9")
doc.add_paragraph(r"= 3180 - 540")
doc.add_paragraph(r"= 2640 kNm")
doc.add_paragraph("Calculation of Shear Force:")
doc.add_paragraph(r"V_u = \alpha f_ck b_w d + 0.8 f_yk A_s \sin \alpha")
doc.add_paragraph(r"= 0.7 \cdot 420 \cdot 833.018 \cdot 793.3 + 300 \cdot 968.2 \cdot \sin 45")
doc.add_paragraph(r"= 71825772.30 + 164410.8")
doc.add_paragraph(r"= 71990183.1")

# 3.2 Calculating the Bent-Up
doc.add_heading('3.2 Calculating the Bent-Up', level=2)
doc.add_paragraph(r"V_u = V_cs")
doc.add_paragraph(r"V_cs = \alpha f_{ck} b_{w} d + f_{yk} \frac{A_{s}}{s} h_{0}")
doc.add_paragraph(r"= 0.7 \cdot 420 \cdot 833.018 \cdot 793.3 + 300 \cdot \frac{200 \cdot 200}{200}")
doc.add_paragraph(r"= 71825772.30")
doc.add_paragraph(r"V_s0 = 0.8 f_{yk} A_{s} \sin \alpha")
doc.add_paragraph(r"= 0.8 \cdot 360 \cdot 968.2 \cdot \sin 45")
doc.add_paragraph(r"= 164410.8")
doc.add_paragraph(r"V_u = \alpha f_{ck} b_{w} d + 0.8 f_{yk} A_{s} \sin \alpha")
doc.add_paragraph(r"= 71825772.30 + 164410.8")
doc.add_paragraph(r"= 71990183.1")

# 3.3 Calculating the Load Bearing Capacity of the Oblique Section
doc.add_heading('3.3 Calculating the Load Bearing Capacity of the Oblique Section', level=2)
doc.add_paragraph("Checking Size and Bearing Capacity of Normal Section and Oblique Section:")
doc.add_paragraph(r"\psi = (b_f - b_w) h_f")
doc.add_paragraph(r"= (980 - 310) \cdot 30.4")
doc.add_paragraph(r"= 0.524")
doc.add_paragraph(r"P_{ec} = A_s = 293 \cdot \alpha")
doc.add_paragraph(r"= 0.5 b_w (b_f - b_w) h_f")
doc.add_paragraph(r"= 0.5 \cdot 310 \cdot 30.4")
doc.add_paragraph(r"= 0.034")

# 4.1 Checking the Size and Bearing Capacity of Normal Section and Oblique Section
doc.add_heading('4.1 Checking the Size and Bearing Capacity of Normal Section and Oblique Section', level=2)
doc.add_paragraph("Checking Size and Bearing Capacity of Normal Section and Oblique Section:")
doc.add_paragraph(r"\psi = (b_f - b_w) h_f")
doc.add_paragraph(r"= (980 - 310) \cdot 30.4")
doc.add_paragraph(r"= 0.524")
doc.add_paragraph(r"P_{ec} = A_s = 293 \cdot \alpha")
doc.add_paragraph(r"= 0.5 b_w (b_f - b_w) h_f")
doc.add_paragraph(r"= 0.5 \cdot 310 \cdot 30.4")
doc.add_paragraph(r"= 0.034")

# 5.1 Calculating the Crack Width
doc.add_heading('5.1 Calculating the Crack Width', level=2)
doc.add_paragraph("Width of Cracks:")
doc.add_paragraph(r"w_m = \frac{A \cdot \psi_s}{E_s} \cdot (1.1 - \frac{d_y}{E_c})")
doc.add_paragraph(r"= \frac{200 \cdot 10^{-3} \cdot 0.85}{30000} \cdot (1.1 - 0.65 \cdot 0.33)")
doc.add_paragraph(r"= 0.179 mm")

# 6.1 Calculating the Deflection
doc.add_heading('6.1 Calculating the Deflection', level=2)
doc.add_paragraph("2) Calculate reaction of roller support at the point B:")
doc.add_paragraph(r"R_B = (q_1 \cdot 6 \cdot (6/2) + q_2 \cdot 2 \cdot (2/2)) / 6 = (60 \cdot 6 \cdot 3 + 150 \cdot 2 \cdot 1) / 6 = 530.00 (kN)")
doc.add_paragraph("3) Calculate reaction of pin support at the point A:")
doc.add_paragraph(r"R_A = (q_1 \cdot 6 \cdot (6/2) + q_2 \cdot 2 \cdot (2/2)) / 6 - R_B = (60 \cdot 6 \cdot 3 + 150 \cdot 2 \cdot 1) / 6 - 530.00 = 130.00 (kN)")
doc.add_paragraph("4) Solve this system of equations:")
doc.add_paragraph("H_A = 0 (kN)")
doc.add_paragraph("5) The sum of the forces about the Oy axis is zero:")
doc.add_paragraph(r"∑_i=1^n P_iy = 0: -q_1 \cdot 6 + q_2 \cdot 2 - R_A - R_B = 0")
doc.add_paragraph("The values of Q at the edges of the span:")
doc.add_paragraph(r"Q_1 (0) = +130.00 - 60 \cdot (0 - 0) = 130 (kN)")
doc.add_paragraph(r"Q_1 (6) = +130.00 - 60 \cdot (6 - 0) = -230 (kN)")
doc.add_paragraph("The value of Q on this span that crosses the horizontal axis. Intersection point: x = 2.17.")
doc.add_paragraph("Determine the equations for the bending moment (M):")
doc.add_paragraph(r"M(x_1) = +R_A \cdot (x_1) - q_1 \cdot (x_1^2 / 2)")

# Save the document
file_path = "/mnt/data/Course_Design_Report.docx"
doc.save(file_path)

file_path
